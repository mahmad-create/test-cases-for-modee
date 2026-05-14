#!/usr/bin/env python3
"""
CI entry point: extract SRS Front End DOCX (if present) and rebuild test artifacts.

Detects a real Word file under Requirements/BRD/ whose name suggests
``SRS Front End … ACC … V1.4`` (ASCII or en-dash), writes
``_extracted_SRS_Front_End_ACC_V1.4.txt``, then runs:

  - scripts/export_srs_v14_tests.py  → template TSV/XLSX
  - scripts/generate_brd_tests_srsfe_v14.py  → BRD-style pack (set BRD_STABLE_OUTPUT=1 in CI)

Extend this script when you add more BRD → generator mappings.
"""
from __future__ import annotations

import os
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
BRD_DIR = ROOT / "Requirements" / "BRD"
EXTRACT_OUT = BRD_DIR / "_extracted_SRS_Front_End_ACC_V1.4.txt"


def _is_zip_docx(path: Path) -> bool:
    try:
        return path.suffix.lower() == ".docx" and path.read_bytes()[:4] == b"PK\x03\x04"
    except OSError:
        return False


def _is_srs_front_end_acc(path: Path) -> bool:
    n = path.name.casefold().replace("\u2013", "-").replace("\u2014", "-")
    return all(k in n for k in ("srs", "front", "end", "acc")) and ("v1" in n or "1.4" in n)


def _find_srs_docx() -> Path | None:
    if not BRD_DIR.is_dir():
        return None
    for p in sorted(BRD_DIR.glob("*.docx")):
        if _is_zip_docx(p) and _is_srs_front_end_acc(p):
            return p
    return None


def _extract_docx(docx: Path, out: Path) -> None:
    from docx import Document

    d = Document(str(docx))
    paras = [t.strip() for t in (x.text for x in d.paragraphs) if t.strip()]
    tbl: list[str] = []
    for t in d.tables:
        for r in t.rows:
            if any(c.text.strip() for c in r.cells):
                tbl.append(" | ".join(c.text.strip() for c in r.cells))
    out.write_text("\n".join(paras + ["\n--- TABLES ---\n"] + tbl), encoding="utf-8")
    print(f"Wrote extract {out.relative_to(ROOT)} ({out.stat().st_size} bytes) from {docx.name}")


def _run(script: str) -> None:
    cmd = [sys.executable, str(ROOT / "scripts" / script)]
    print("+", " ".join(cmd))
    subprocess.check_call(cmd, cwd=str(ROOT), env=os.environ.copy())


def main() -> int:
    srs = _find_srs_docx()
    if srs:
        _extract_docx(srs, EXTRACT_OUT)
    elif not EXTRACT_OUT.exists():
        print("No SRS Front End ACC V1.4 .docx found and no extract on disk; skipping generators.")
        return 0

    _run("export_srs_v14_tests.py")
    _run("generate_brd_tests_srsfe_v14.py")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
